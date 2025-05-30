"""
Document processing service for extracting text and creating embeddings.
"""

import os
import uuid
import logging
from typing import List, Optional
from pathlib import Path
import aiofiles
from fastapi import UploadFile, HTTPException

# Document processing imports
import pypdf
from docx import Document as DocxDocument
import chromadb
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
import openai

from app.utils.config import settings

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing service."""

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)

        # Initialize ChromaDB
        try:
            self.chroma_client = chromadb.PersistentClient(path=settings.CHROMA_DB_PATH)
            self.collection = self.chroma_client.get_or_create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
            logger.info("ChromaDB initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize ChromaDB: {e}")
            raise

        # Initialize OpenAI embeddings
        try:
            if settings.OPENAI_API_KEY:
                self.embeddings = OpenAIEmbeddings(
                    openai_api_key=settings.OPENAI_API_KEY,
                    model="text-embedding-3-small"  # More cost-effective than ada-002
                )
                self.use_embeddings = True
                logger.info("OpenAI embeddings initialized successfully")
            else:
                logger.warning("No OpenAI API key provided. Document search will be limited.")
                self.embeddings = None
                self.use_embeddings = False
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI embeddings: {e}")
            self.embeddings = None
            self.use_embeddings = False

        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    async def save_file(self, file: UploadFile) -> tuple[str, str]:
        """Save uploaded file to disk."""
        # Validate file type
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in settings.ALLOWED_FILE_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_extension} not allowed"
            )
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.upload_dir / unique_filename
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            if len(content) > settings.MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail="File too large"
                )
            await f.write(content)
        
        return str(file_path), unique_filename
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                return self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting text: {str(e)}"
            )
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text += f"Page {page_num + 1}:\n{page_text}\n\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue

            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {e}")
            raise
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX file")

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {e}")
            raise

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        if text.strip():
                            return text.strip()
                except UnicodeDecodeError:
                    continue

            raise ValueError("Could not decode the text file with any supported encoding")
        except Exception as e:
            logger.error(f"Error extracting TXT text from {file_path}: {e}")
            raise
    
    async def process_document(self, file_path: str, document_id: int) -> bool:
        """Process document and store embeddings."""
        try:
            logger.info(f"Starting to process document {document_id}: {file_path}")

            # Extract text
            text = self.extract_text(file_path)
            logger.info(f"Extracted {len(text)} characters from document {document_id}")

            if not text.strip():
                logger.warning(f"No text extracted from document {document_id}")
                return False

            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            logger.info(f"Split document {document_id} into {len(chunks)} chunks")

            if not self.use_embeddings:
                logger.warning("OpenAI embeddings not available. Storing text without embeddings.")
                # Store chunks without embeddings for basic text search
                for i, chunk in enumerate(chunks):
                    chunk_id = f"doc_{document_id}_chunk_{i}"
                    self.collection.add(
                        documents=[chunk],
                        metadatas=[{
                            "document_id": document_id,
                            "chunk_index": i,
                            "file_path": file_path
                        }],
                        ids=[chunk_id]
                    )
                return True

            # Create embeddings and store in ChromaDB
            for i, chunk in enumerate(chunks):
                try:
                    chunk_id = f"doc_{document_id}_chunk_{i}"

                    # Generate embedding using OpenAI
                    embedding = await self._generate_embedding(chunk)

                    self.collection.add(
                        embeddings=[embedding],
                        documents=[chunk],
                        metadatas=[{
                            "document_id": document_id,
                            "chunk_index": i,
                            "file_path": file_path,
                            "chunk_length": len(chunk)
                        }],
                        ids=[chunk_id]
                    )

                    logger.debug(f"Processed chunk {i+1}/{len(chunks)} for document {document_id}")

                except Exception as e:
                    logger.error(f"Error processing chunk {i} of document {document_id}: {e}")
                    continue

            logger.info(f"Successfully processed document {document_id}")
            return True

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            return False

    async def _generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using OpenAI."""
        try:
            # Use the embed_query method which is async-compatible
            embedding = self.embeddings.embed_query(text)
            return embedding
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            raise
    
    async def search_documents(self, query: str, n_results: int = 5) -> List[dict]:
        """Search for relevant document chunks."""
        try:
            logger.info(f"Searching documents for query: {query[:100]}...")

            if self.use_embeddings and self.embeddings:
                # Use semantic search with embeddings
                query_embedding = await self._generate_embedding(query)
                results = self.collection.query(
                    query_embeddings=[query_embedding],
                    n_results=n_results
                )
                logger.info(f"Found {len(results['documents'][0]) if results['documents'] else 0} results using embeddings")
            else:
                # Fallback to text-based search
                logger.info("Using text-based search (no embeddings available)")
                results = self.collection.query(
                    query_texts=[query],
                    n_results=n_results
                )
                logger.info(f"Found {len(results['documents'][0]) if results['documents'] else 0} results using text search")

            if not results["documents"] or not results["documents"][0]:
                logger.info("No documents found for query")
                return []

            search_results = []
            for i in range(len(results["documents"][0])):
                result = {
                    "content": results["documents"][0][i],
                    "metadata": results["metadatas"][0][i],
                }

                # Add distance/score if available
                if results.get("distances") and results["distances"][0]:
                    result["distance"] = results["distances"][0][i]
                    result["similarity"] = 1 - results["distances"][0][i]  # Convert distance to similarity

                search_results.append(result)

            logger.info(f"Returning {len(search_results)} search results")
            return search_results

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []

    def get_document_stats(self) -> dict:
        """Get statistics about stored documents."""
        try:
            count = self.collection.count()
            return {
                "total_chunks": count,
                "embeddings_enabled": self.use_embeddings,
                "collection_name": self.collection.name
            }
        except Exception as e:
            logger.error(f"Error getting document stats: {e}")
            return {"error": str(e)}

    def delete_document_chunks(self, document_id: int) -> bool:
        """Delete all chunks for a specific document."""
        try:
            # Get all chunk IDs for this document
            results = self.collection.get(
                where={"document_id": document_id}
            )

            if results["ids"]:
                self.collection.delete(ids=results["ids"])
                logger.info(f"Deleted {len(results['ids'])} chunks for document {document_id}")
                return True
            else:
                logger.info(f"No chunks found for document {document_id}")
                return True

        except Exception as e:
            logger.error(f"Error deleting chunks for document {document_id}: {e}")
            return False


# Global instance
document_processor = DocumentProcessor()
